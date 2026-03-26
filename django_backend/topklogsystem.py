import os

os.environ["ANONYMIZED_TELEMETRY"] = "false"
os.environ["DISABLE_TELEMETRY"] = "1"
os.environ["CHROMA_TELEMETRY_ENABLED"] = "false"

import json
import logging
import threading
import pandas as pd
from typing import Any, Dict, List, Optional

# langchain
from langchain.prompts import (
    ChatPromptTemplate,
    HumanMessagePromptTemplate,
    SystemMessagePromptTemplate,
    MessagesPlaceholder,
)
from langchain_core.messages import (
    HumanMessage,
    AIMessage,
)
from langchain_ollama import OllamaLLM, OllamaEmbeddings

# llama-index & chroma
import chromadb
from llama_index.core import Settings
from llama_index.core import Document
from llama_index.core import VectorStoreIndex, StorageContext
from llama_index.vector_stores.chroma import ChromaVectorStore

# 日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

import re
from docx import Document as DocxDocument
from PyPDF2 import PdfReader


class TopKLogSystem:
    def __init__(
        self,
        log_path: str,
        llm: str,
        embedding_model: str,
    ) -> None:
        self.embedding_model = OllamaEmbeddings(model=embedding_model)
        self._default_llm_name = llm
        self._llm_cache: Dict[str, OllamaLLM] = {}
        self._llm_lock = threading.RLock()
        self.llm = self._get_or_create_llm(llm)

        Settings.llm = self.llm
        Settings.embed_model = self.embedding_model
        Settings.chunk_size = 400
        Settings.chunk_overlap = 40

        self.log_path = log_path
        self.log_index = None
        self.vector_store = None
        self._build_vectorstore()

    def _build_vectorstore(self):
        vector_store_path = "./data/vector_stores"

        if os.path.exists(vector_store_path):
            logger.info(f"加载现有向量数据库索引: {vector_store_path}")
            chroma_client = chromadb.PersistentClient(path=vector_store_path)
            log_collection = chroma_client.get_collection("log_collection")
            log_vector_store = ChromaVectorStore(chroma_collection=log_collection)

            self.log_index = VectorStoreIndex.from_vector_store(
                vector_store=log_vector_store
            )
            self.vector_store = log_vector_store
            return

        logger.info(f"构建新向量数据库: {vector_store_path}")
        os.makedirs(vector_store_path, exist_ok=True)

        chroma_client = chromadb.PersistentClient(path=vector_store_path)
        log_collection = chroma_client.get_or_create_collection("log_collection")

        log_vector_store = ChromaVectorStore(chroma_collection=log_collection)
        self.vector_store = log_vector_store

        log_storage_context = StorageContext.from_defaults(
            vector_store=log_vector_store
        )
        
        log_documents = self._load_documents(self.log_path)
        if log_documents:
            self.log_index = VectorStoreIndex.from_documents(
                log_documents,
                storage_context=log_storage_context,
                show_progress=True,
            )
            logger.info(f"日志库索引构建完成，共 {len(log_documents)} 条数据")

    @staticmethod
    def _load_documents(data_path: str) -> List[Document]:
        if not os.path.exists(data_path):
            return []
        documents = []
        for root, _, files in os.walk(data_path):
            for file in files:
                ext = os.path.splitext(file)[1]
                if ext not in [
                    ".txt", ".md", ".json", ".jsonl", ".csv",
                    ".log", ".xml", ".yaml", ".yml", ".docx", ".pdf",
                ]:
                    continue
                
                file_path = os.path.join(root, file)
                
                if ext == ".csv":
                    documents.extend(TopKLogSystem._process_csv(file_path))
                elif ext in [".json", ".jsonl"]:
                    documents.extend(TopKLogSystem._process_json(file_path, ext))
                elif ext in [".yaml", ".yml"]:
                    documents.extend(TopKLogSystem._process_yaml(file_path))
                elif ext == ".xml":
                    documents.extend(TopKLogSystem._process_xml(file_path))
                elif ext == ".log":
                    documents.extend(TopKLogSystem._process_log(file_path))
                elif ext == ".docx":
                    documents.extend(TopKLogSystem._process_docx(file_path))
                elif ext == ".pdf":
                    documents.extend(TopKLogSystem._process_pdf(file_path))
                else:
                    documents.extend(TopKLogSystem._process_text(file_path))
                    
        return documents

    @staticmethod
    def _process_csv(file_path: str) -> List[Document]:
        documents = []
        for chunk in pd.read_csv(file_path, chunksize=1000, on_bad_lines="skip"):
            for row in chunk.itertuples(index=False):
                content = str(row).replace("Pandas", " ")
                documents.append(Document(text=content))
        return documents

    @staticmethod
    def _normalize_risk_level(raw_value: Any) -> str:
        lowered = str(raw_value or "").strip().lower()
        mapping = {
            "critical": "Critical",
            "high": "High",
            "medium": "Medium",
            "low": "Low",
            "info": "Info",
        }
        return mapping.get(lowered, "Info")

    @staticmethod
    def _normalize_tags(raw_tags: Any) -> List[str]:
        if raw_tags is None:
            return []
        if isinstance(raw_tags, list):
            return [str(tag).strip() for tag in raw_tags if str(tag).strip()]
        if isinstance(raw_tags, str):
            return [tag.strip() for tag in raw_tags.split(",") if tag.strip()]
        return []

    @staticmethod
    def _normalize_mitre_ids(raw_ids: Any) -> List[str]:
        if raw_ids is None:
            return []
        if isinstance(raw_ids, list):
            return [str(item).strip() for item in raw_ids if str(item).strip()]
        if isinstance(raw_ids, str):
            extracted = re.findall(r"T\d{4}(?:\.\d{3})?", raw_ids.upper())
            return sorted(set(extracted))
        return []

    @staticmethod
    def _serialize_metadata_list(values: List[str]) -> Optional[str]:
        cleaned_values = [str(item).strip() for item in values if str(item).strip()]
        if not cleaned_values:
            return None
        return ",".join(cleaned_values)

    @staticmethod
    def _to_int(raw_value: Any) -> Optional[int]:
        if raw_value is None:
            return None
        if isinstance(raw_value, bool):
            return int(raw_value)
        if isinstance(raw_value, (int, float)):
            return int(raw_value)
        text = str(raw_value).strip()
        if not text:
            return None
        if text.isdigit():
            return int(text)
        return None

    @staticmethod
    def _build_structured_doc(
        record: Dict[str, Any], file_path: str, line_no: Optional[int] = None
    ) -> Optional[Document]:
        search_content = str(record.get("search_content") or "").strip()
        if not search_content:
            return None

        metadata: Dict[str, Any] = {
            "_id": record.get("_id"),
            "db_type": record.get("db_type"),
            "risk_level": TopKLogSystem._normalize_risk_level(record.get("risk_level")),
            "cve_id": record.get("cve_id"),
            "ioc_value": record.get("ioc_value"),
            "source": record.get("source"),
            "confidence": TopKLogSystem._to_int(record.get("confidence")),
            "raw_content_hash": record.get("raw_content_hash"),
            "verified": int(bool(record.get("verified", False))),
            "source_dataset": record.get("source_dataset"),
            "source_priority": TopKLogSystem._to_int(record.get("source_priority")),
            "mitre_attack_id": TopKLogSystem._serialize_metadata_list(
                TopKLogSystem._normalize_mitre_ids(record.get("mitre_attack_id"))
            ),
            "tags": TopKLogSystem._serialize_metadata_list(
                TopKLogSystem._normalize_tags(record.get("tags"))
            ),
        }
        metadata = {k: v for k, v in metadata.items() if v not in (None, "", [], {})}
        metadata["record_file"] = os.path.basename(file_path)
        if line_no is not None:
            metadata["record_line"] = line_no

        return Document(text=search_content, metadata=metadata)

    @staticmethod
    def _process_json(file_path: str, ext: str) -> List[Document]:
        documents: List[Document] = []
        with open(file_path, "r", encoding="utf-8") as f:
            if ext == ".json":
                data = json.load(f)
                if isinstance(data, list):
                    for item in data:
                        if isinstance(item, dict):
                            doc = TopKLogSystem._build_structured_doc(item, file_path)
                            if doc is not None:
                                documents.append(doc)
                elif isinstance(data, dict):
                    doc = TopKLogSystem._build_structured_doc(data, file_path)
                    if doc is not None:
                        documents.append(doc)
            elif ext == ".jsonl":
                for line_no, line in enumerate(f, start=1):
                    stripped = line.strip()
                    if not stripped:
                        continue
                    data = json.loads(stripped)
                    if isinstance(data, dict):
                        doc = TopKLogSystem._build_structured_doc(
                            data,
                            file_path,
                            line_no=line_no,
                        )
                        if doc is not None:
                            documents.append(doc)
        return documents

    @staticmethod
    def _process_yaml(file_path: str) -> List[Document]:
        documents = []
        with open(file_path, "r", encoding="utf-8") as f:
            documents.append(Document(text=f.read()))
        return documents

    @staticmethod
    def _process_xml(file_path: str) -> List[Document]:
        documents = []
        with open(file_path, "r", encoding="utf-8") as f:
            documents.append(Document(text=f.read()))
        return documents

    @staticmethod
    def _process_log(file_path: str) -> List[Document]:
        documents = []
        with open(file_path, "r", encoding="utf-8") as f:
            for line in f:
                documents.append(Document(text=line.strip()))
        return documents

    @staticmethod
    def _process_docx(file_path: str) -> List[Document]:
        documents = []
        doc = DocxDocument(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                documents.append(Document(text=paragraph.text.strip()))
        return documents

    @staticmethod
    def _process_pdf(file_path: str) -> List[Document]:
        documents = []
        reader = PdfReader(file_path)
        for page in reader.pages:
            text = page.extract_text()
            if text:
                documents.append(Document(text=text.strip()))
        return documents

    @staticmethod
    def _process_text(file_path: str) -> List[Document]:
        documents = []
        with open(file_path, "r", encoding="utf-8") as f:
            documents.append(Document(text=f.read()))
        return documents

    @staticmethod
    def _normalize_source_priority(source: str, source_priority: Any) -> float:
        if isinstance(source_priority, (int, float)):
            return max(0.0, min(float(source_priority), 100.0)) / 100.0

        source_rank = {
            "nvd": 0.95,
            "cisa": 0.92,
            "threatfox": 0.85,
            "urlhaus": 0.8,
            "otx": 0.78,
        }
        return source_rank.get(str(source or "").strip().lower(), 0.6)

    @staticmethod
    def _build_query_terms(query: str) -> List[str]:
        lowered = query.lower()
        tokens = set(re.findall(r"[a-zA-Z0-9\-_.]+", lowered))
        tokens.update(
            match.lower()
            for match in re.findall(r"CVE-\d{4}-\d+", query, flags=re.IGNORECASE)
        )
        tokens.update(
            match.lower()
            for match in re.findall(r"T\d{4}(?:\.\d{3})?", query, flags=re.IGNORECASE)
        )
        return sorted(token for token in tokens if len(token) >= 3)

    @staticmethod
    def _matches_filters(metadata: Dict[str, Any], filters: Optional[Dict[str, Any]]) -> bool:
        if not filters:
            return True

        for key, expected in filters.items():
            actual = metadata.get(key)
            if expected is None:
                continue

            if isinstance(expected, list):
                expected_set = {str(item).strip().lower() for item in expected if str(item).strip()}
                if not expected_set:
                    continue

                if isinstance(actual, list):
                    actual_set = {str(item).strip().lower() for item in actual if str(item).strip()}
                elif isinstance(actual, str):
                    actual_set = {
                        item.strip().lower()
                        for item in re.split(r"[;,]", actual)
                        if item.strip()
                    }
                else:
                    actual_set = {str(actual).strip().lower()} if actual is not None else set()

                if expected_set.isdisjoint(actual_set):
                    return False
            else:
                expected_text = str(expected).strip().lower()
                actual_text = str(actual).strip().lower() if actual is not None else ""
                if expected_text and actual_text != expected_text:
                    return False

        return True

    @staticmethod
    def _extract_query_signals(query: str) -> Dict[str, List[str]]:
        cve_ids = sorted({match.upper() for match in re.findall(r"CVE-\d{4}-\d+", query, flags=re.IGNORECASE)})
        ip_addresses = sorted({match for match in re.findall(r"\b\d{1,3}(?:\.\d{1,3}){3}\b", query)})
        hashes = sorted({match.lower() for match in re.findall(r"\b[a-fA-F0-9]{32}\b|\b[a-fA-F0-9]{40}\b|\b[a-fA-F0-9]{64}\b", query)})
        mitre_ids = sorted({match.upper() for match in re.findall(r"T\d{4}(?:\.\d{3})?", query, flags=re.IGNORECASE)})

        exact_terms = sorted({*{item.lower() for item in cve_ids}, *{item.lower() for item in ip_addresses}, *hashes, *{item.lower() for item in mitre_ids}})

        return {
            "cve_ids": cve_ids,
            "ip_addresses": ip_addresses,
            "hashes": hashes,
            "mitre_ids": mitre_ids,
            "exact_terms": exact_terms,
        }

    @staticmethod
    def _deserialize_metadata_list(raw_value: Any) -> List[str]:
        if raw_value is None:
            return []
        if isinstance(raw_value, list):
            return [str(item).strip() for item in raw_value if str(item).strip()]
        if isinstance(raw_value, str):
            return [item.strip() for item in re.split(r"[;,]", raw_value) if item.strip()]
        return []

    @staticmethod
    def _build_searchable_text(text: str, metadata: Dict[str, Any]) -> str:
        parts = [text]
        for key in (
            "db_type",
            "risk_level",
            "cve_id",
            "ioc_value",
            "source",
            "raw_content_hash",
            "record_file",
            "record_line",
            "mitre_attack_id",
            "tags",
        ):
            value = metadata.get(key)
            if value is not None:
                parts.append(str(value))
        return " ".join(parts).lower()

    def _score_keyword_match(
        self,
        text: str,
        metadata: Dict[str, Any],
        query_terms: List[str],
        exact_terms: List[str],
    ) -> float:
        searchable_text = self._build_searchable_text(text, metadata)

        exact_hit_count = 0
        exact_metadata_values = [
            str(metadata.get("cve_id") or "").strip().lower(),
            str(metadata.get("ioc_value") or "").strip().lower(),
            str(metadata.get("raw_content_hash") or "").strip().lower(),
            str(metadata.get("source") or "").strip().lower(),
        ]
        exact_metadata_values.extend(item.lower() for item in self._deserialize_metadata_list(metadata.get("mitre_attack_id")))

        for term in exact_terms:
            if not term:
                continue
            if term in searchable_text or term in exact_metadata_values:
                exact_hit_count += 1

        token_hits = sum(1 for term in query_terms if term and term in searchable_text)

        if exact_hit_count:
            return min(1.0, 0.72 + exact_hit_count * 0.12 + min(token_hits, 4) * 0.04)
        if token_hits == 0:
            return 0.0
        return min(0.65, 0.18 * token_hits)

    def _collect_candidates(
        self,
        query: str,
        top_k: int,
        use_keyword: bool,
        filters: Optional[Dict[str, Any]],
        vector_floor: float,
        keyword_floor: float,
    ) -> tuple[Dict[str, Dict[str, Any]], Dict[str, List[str]]]:
        query_signals = self._extract_query_signals(query)
        query_terms = self._build_query_terms(query)
        candidate_map: Dict[str, Dict[str, Any]] = {}

        def _candidate_key(text: str, metadata: Dict[str, Any]) -> str:
            return str(metadata.get("raw_content_hash") or metadata.get("_id") or text)

        # 1. 意图短路：直接利用 Chroma DB 原生查询提取的强特征 (如 CVE_ID)
        if self.vector_store and hasattr(self.vector_store, "_collection"):
            collection = self.vector_store._collection
            
            # 针对识别到的 CVE IDs 强制精确召回
            for cve in query_signals.get("cve_ids", []):
                exact_hits = collection.get(where={"cve_id": cve.upper()})
                if exact_hits and exact_hits.get("ids"):
                    for i in range(len(exact_hits["ids"])):
                        text = exact_hits["documents"][i]
                        metadata = exact_hits["metadatas"][i]
                        
                        if not self._matches_filters(metadata, filters):
                            continue
                            
                        key = _candidate_key(text, metadata)
                        item = candidate_map.setdefault(
                            key,
                            {
                                "content": text,
                                "metadata": metadata,
                                "vector_score": 1.0,  # 强匹配直接赋满分
                                "keyword_score": 1.0,
                                "channels": set(),
                            },
                        )
                        item["channels"].add("exact_metadata")

        # 2. 常规向量召回
        retriever = self.log_index.as_retriever(similarity_top_k=max(top_k * 4, 20))
        vector_hits = retriever.retrieve(query)

        for node in vector_hits:
            text = (getattr(node, "text", "") or "").strip()
            if not text:
                continue

            metadata = dict(getattr(node, "metadata", {}) or {})
            if not self._matches_filters(metadata, filters):
                continue

            vector_score = float(getattr(node, "score", 0.0) or 0.0)
            if vector_score < vector_floor:
                continue

            key = _candidate_key(text, metadata)
            item = candidate_map.setdefault(
                key,
                {
                    "content": text,
                    "metadata": metadata,
                    "vector_score": 0.0,
                    "keyword_score": 0.0,
                    "channels": set(),
                },
            )
            item["vector_score"] = max(item["vector_score"], vector_score)
            item["channels"].add("vector")

        # 3. 修复后的关键词得分计算
        # 仅对已进入 candidate_map 的候选者进行关键词算分，抛弃失效的空 docstore 遍历
        if use_keyword:
            for key, item in candidate_map.items():
                text = item["content"]
                metadata = item["metadata"]
                
                keyword_score = self._score_keyword_match(
                    text, metadata, query_terms, query_signals["exact_terms"]
                )
                
                # 若来自 exact_metadata 强制召回，保留其 1.0 的满分
                if "exact_metadata" not in item["channels"]:
                    if keyword_score < keyword_floor:
                        # 仅依靠向量召回但关键词极低的可以考虑剔除，此处保留原始逻辑仅更新分数
                        item["keyword_score"] = max(item["keyword_score"], keyword_score)
                    else:
                        item["keyword_score"] = max(item["keyword_score"], keyword_score)
                        item["channels"].add("keyword")

        return candidate_map, query_signals

    @staticmethod
    def _rank_candidates(
        candidate_map: Dict[str, Dict[str, Any]],
        query_signals: Dict[str, List[str]],
    ) -> List[Dict[str, Any]]:
        risk_weight = {
            "Critical": 1.0,
            "High": 0.85,
            "Medium": 0.65,
            "Low": 0.4,
            "Info": 0.25,
        }
        ranked_items: List[Dict[str, Any]] = []

        for item in candidate_map.values():
            metadata = item["metadata"]
            searchable_text = TopKLogSystem._build_searchable_text(item["content"], metadata)

            exact_match_boost = 0.0
            for exact_term in query_signals.get("exact_terms", []):
                if exact_term and exact_term in searchable_text:
                    exact_match_boost = 1.0
                    break

            confidence = metadata.get("confidence")
            confidence_norm = 0.0
            if isinstance(confidence, (int, float)):
                confidence_norm = max(0.0, min(float(confidence), 100.0)) / 100.0

            source_priority = TopKLogSystem._normalize_source_priority(
                str(metadata.get("source", "")),
                metadata.get("source_priority"),
            )
            risk_level = TopKLogSystem._normalize_risk_level(metadata.get("risk_level"))
            risk_score = risk_weight.get(risk_level, 0.25)

            final_score = (
                item["vector_score"] * 0.5
                + item["keyword_score"] * 0.2
                + confidence_norm * 0.15
                + source_priority * 0.1
                + risk_score * 0.05
                + exact_match_boost
            )

            ranked_items.append(
                {
                    "content": item["content"],
                    "score": round(final_score, 6),
                    "source": "+".join(sorted(item["channels"])) or "hybrid",
                    "metadata": metadata,
                    "evidence": {
                        "db_type": metadata.get("db_type"),
                        "risk_level": risk_level,
                        "source": metadata.get("source"),
                        "confidence": metadata.get("confidence"),
                        "raw_content_hash": metadata.get("raw_content_hash"),
                    },
                    "vector_score": round(float(item["vector_score"]), 6),
                    "keyword_score": round(float(item["keyword_score"]), 6),
                    "exact_match_boost": exact_match_boost,
                }
            )

        ranked_items.sort(key=lambda x: x["score"], reverse=True)
        return ranked_items

    @staticmethod
    def _select_group_key(item: Dict[str, Any]) -> tuple[str, str]:
        metadata = item.get("metadata", {}) or {}
        cve_id = str(metadata.get("cve_id") or "").strip().upper()
        if cve_id:
            return f"cve:{cve_id}", "cve"

        ioc_value = str(metadata.get("ioc_value") or "").strip().lower()
        if ioc_value:
            return f"ioc:{ioc_value}", "ioc"

        raw_content_hash = str(metadata.get("raw_content_hash") or "").strip().lower()
        if raw_content_hash:
            return f"hash:{raw_content_hash}", "raw_content_hash"

        tags = TopKLogSystem._deserialize_metadata_list(metadata.get("tags"))
        if tags:
            return f"attack_type:{tags[0].strip().lower()}", "attack_type"

        db_type = str(metadata.get("db_type") or "unknown").strip().lower()
        source = str(metadata.get("source") or "unknown").strip().lower()
        return f"db:{db_type}:{source}", "db_type"

    @staticmethod
    def _build_prompt_metadata(metadata: Dict[str, Any]) -> Dict[str, Any]:
        return {
            "_id": metadata.get("_id"),
            "db_type": metadata.get("db_type"),
            "risk_level": metadata.get("risk_level"),
            "cve_id": metadata.get("cve_id"),
            "ioc_value": metadata.get("ioc_value"),
            "source": metadata.get("source"),
            "confidence": metadata.get("confidence"),
            "raw_content_hash": metadata.get("raw_content_hash"),
            "source_dataset": metadata.get("source_dataset"),
            "source_priority": metadata.get("source_priority"),
            "verified": bool(metadata.get("verified", False)),
            "record_file": metadata.get("record_file"),
            "record_line": metadata.get("record_line"),
            "mitre_attack_id": TopKLogSystem._deserialize_metadata_list(metadata.get("mitre_attack_id")),
            "tags": TopKLogSystem._deserialize_metadata_list(metadata.get("tags")),
        }

    def _group_retrieval_results(self, ranked_items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        grouped: Dict[str, Dict[str, Any]] = {}

        for item in ranked_items:
            group_key, group_type = self._select_group_key(item)
            metadata = item.get("metadata", {}) or {}
            prompt_metadata = self._build_prompt_metadata(metadata)
            tags = prompt_metadata.get("tags", [])
            mitre_attack_id = prompt_metadata.get("mitre_attack_id", [])

            group = grouped.setdefault(
                group_key,
                {
                    "group_key": group_key,
                    "group_type": group_type,
                    "score": 0.0,
                    "source": set(),
                    "members": [],
                    "entity_summary": {
                        "cve_id": prompt_metadata.get("cve_id"),
                        "ioc_values": [],
                        "raw_content_hashes": [],
                        "attack_types": [],
                        "mitre_attack_id": [],
                    },
                },
            )

            group["score"] = max(group["score"], float(item.get("score", 0.0) or 0.0))
            group["source"].add(item.get("source", "unknown"))
            group["members"].append(
                {
                    "content": item.get("content", ""),
                    "score": item.get("score", 0.0),
                    "source": item.get("source", "unknown"),
                    "vector_score": item.get("vector_score", 0.0),
                    "keyword_score": item.get("keyword_score", 0.0),
                    "exact_match_boost": item.get("exact_match_boost", 0.0),
                    "metadata": prompt_metadata,
                    "evidence": item.get("evidence", {}),
                }
            )

            if prompt_metadata.get("ioc_value"):
                ioc_value = prompt_metadata["ioc_value"]
                if ioc_value not in group["entity_summary"]["ioc_values"]:
                    group["entity_summary"]["ioc_values"].append(ioc_value)

            raw_content_hash = prompt_metadata.get("raw_content_hash")
            if raw_content_hash and raw_content_hash not in group["entity_summary"]["raw_content_hashes"]:
                group["entity_summary"]["raw_content_hashes"].append(raw_content_hash)

            for tag in tags:
                if tag not in group["entity_summary"]["attack_types"]:
                    group["entity_summary"]["attack_types"].append(tag)

            for attack_id in mitre_attack_id:
                if attack_id not in group["entity_summary"]["mitre_attack_id"]:
                    group["entity_summary"]["mitre_attack_id"].append(attack_id)

        grouped_items: List[Dict[str, Any]] = []
        for group in grouped.values():
            group["members"].sort(key=lambda x: x.get("score", 0.0), reverse=True)
            representative = group["members"][0] if group["members"] else {}
            grouped_items.append(
                {
                    "group_key": group["group_key"],
                    "group_type": group["group_type"],
                    "score": round(float(group["score"]), 6),
                    "source": "+".join(sorted(group["source"])) or "hybrid",
                    "content": representative.get("content", ""),
                    "metadata": representative.get("metadata", {}),
                    "evidence": representative.get("evidence", {}),
                    "member_count": len(group["members"]),
                    "entity_summary": group["entity_summary"],
                    "evidence_chain": group["members"],
                }
            )

        grouped_items.sort(key=lambda x: x["score"], reverse=True)
        return grouped_items

    def retrieve(
        self,
        query: str,
        top_k: int = 10,
        use_keyword: bool = True,
        filters: Optional[Dict[str, Any]] = None,
    ) -> List[Dict[str, Any]]:
        """
        统一检索接口：意图短路精准识别 -> 召回(向量+关键词) -> 重排(结构化字段加权) -> 证据链聚合 -> 去重(raw_content_hash)。
        """
        if not self.log_index:
            return []
        primary_candidates, query_signals = self._collect_candidates(
            query=query,
            top_k=top_k,
            use_keyword=use_keyword,
            filters=filters,
            vector_floor=0.2,
            keyword_floor=0.22,
        )
        ranked_items = self._rank_candidates(primary_candidates, query_signals)

        best_score = ranked_items[0]["score"] if ranked_items else 0.0
        if not ranked_items or best_score < 0.35:
            relaxed_candidates, relaxed_signals = self._collect_candidates(
                query=query,
                top_k=max(top_k * 8, 40),
                use_keyword=use_keyword,
                filters=filters,
                vector_floor=0.0,
                keyword_floor=0.08,
            )
            relaxed_ranked_items = self._rank_candidates(relaxed_candidates, relaxed_signals)

            relaxed_best_score = relaxed_ranked_items[0]["score"] if relaxed_ranked_items else 0.0
            if relaxed_ranked_items and (
                not ranked_items
                or relaxed_best_score > best_score
                or len(relaxed_ranked_items) > len(ranked_items)
            ):
                ranked_items = relaxed_ranked_items

            if not ranked_items and filters:
                fallback_candidates, fallback_signals = self._collect_candidates(
                    query=query,
                    top_k=max(top_k * 8, 40),
                    use_keyword=use_keyword,
                    filters=None,
                    vector_floor=0.0,
                    keyword_floor=0.08,
                )
                fallback_ranked_items = self._rank_candidates(fallback_candidates, fallback_signals)
                if fallback_ranked_items:
                    ranked_items = fallback_ranked_items

        grouped_items = self._group_retrieval_results(ranked_items)
        return grouped_items[:top_k]

    def retrieve_logs(
        self,
        query: str,
        top_k: int = 10,
        use_keyword: bool = True,
        filter_func=None,
    ) -> List[Dict]:
        all_results = self.retrieve(query, top_k=top_k, use_keyword=use_keyword)
        if filter_func:
            all_results = [item for item in all_results if filter_func(item)]
        return all_results[:top_k]

    def _get_or_create_llm(self, model_name: Optional[str]) -> OllamaLLM:
        target_name = (model_name or self._default_llm_name or "").strip()
        if not target_name:
            target_name = self._default_llm_name

        cached_llm = self._llm_cache.get(target_name)
        if cached_llm is not None:
            return cached_llm

        new_llm = OllamaLLM(model=target_name, temperature=0.1, keep_alive="1h")
        self._llm_cache[target_name] = new_llm
        return new_llm

    def generate_response(
        self,
        query: str,
        context: Dict,
        history: List[Dict] = None,
        model_name: Optional[str] = None,
    ):
        prompt_messages = self._build_prompt(query, context, history)
        llm_to_use = self._get_or_create_llm(model_name)

        with self._llm_lock:
            self.llm = llm_to_use
            Settings.llm = llm_to_use

        for chunk in llm_to_use.stream(prompt_messages):
            yield chunk

    def _build_prompt(
        self, query: str, context: Dict, history: List[Dict] = None
    ) -> List:
        log_data = context.get("log_context", [])
        web_data = context.get("web_context", [])

        system_message = SystemMessagePromptTemplate.from_template(
            """
            你是一个多任务SRE助手。你的首要任务是 **[判断意图]**，然后根据意图选择正确的 **[响应模式]**。
            你有两种响应模式：
            1.  **[SRE分析模式]**: 当用户的问题与故障排查、日志分析、系统错误相关时使用。
            2.  **[常规对话模式]**: 当用户进行常规闲聊 (如 "你好")、历史回顾 (如 "我刚才问了什么") 或提出与日志无关的问题 (如 "介绍一下天津大学") 时使用。

            **[!!! 可用工具 (SRE模式专用) !!!]**
            你现在有两种工具上下文：
            1.  **[日志数据库 (Log DB)]**: 包含本地的、详细的系统日志。
            2.  **[联网搜索 (Web Search)]**: 包含来自互联网的实时信息。

            你的回答必须遵循以下质量要求：
            1.  **专业严谨**：在 [SRE分析模式] 下，你的分析必须基于上下文（[日志数据库] 和/或 [联网搜索]），严禁凭空猜测。
            2.  **优先使用日志**：如果 [日志数据库] 提供了足够的信息，优先使用它。只有当日志信息不足或用户明确询问需要外部知识时，才使用 [联网搜索]。
            3.  **清晰可读**：使用 Markdown 格式（如列表、代码块、粗体）来组织你的回答。
            4.  **上下文感知**：你必须能够 **自主判断** 是否需要结合 **历史对话** 来理解用户的真实意图或SRE问题。

            **[!!! 绝对指令：输出格式 !!!]**
            1.  你 **必须** 且 **只能** 使用 `<think>...</think>` 标签来包裹你的所有内部思考步骤 (包括意图分析、SRE分析框架等)。
            2.  在 `<think>...</think>` 标签之外，你 **必须** 且 **只能** 输出 **最终的、直接面向用户** 的回复。
            3.  最终回复中 **严禁** 包含 "步骤 1"、"步骤 2"、"意图分析"、"根本原因"、"最终回复草稿" 等任何思考过程的字样。
            """
        )

        log_context_payload = {
            "type": "log_evidence_chain",
            "count": len(log_data),
            "items": log_data,
        }
        if not log_data:
            log_context_payload["note"] = "未从日志数据库检索到相关内容"
        log_context_str = json.dumps(log_context_payload, ensure_ascii=False, indent=2)

        web_context_payload = {
            "type": "web_context",
            "count": len(web_data),
            "items": web_data,
        }
        if not web_data:
            web_context_payload["note"] = "未启用或未从联网搜索检索到相关内容"
        web_context_str = json.dumps(web_context_payload, ensure_ascii=False, indent=2)

        user_message_template = HumanMessagePromptTemplate.from_template(
            "{log_context}\n{web_context}\n\n当前用户问题:\n{query}"
        )

        prompt_template = ChatPromptTemplate.from_messages(
            [
                system_message,
                MessagesPlaceholder(variable_name="chat_history"),
                user_message_template,
            ]
        )

        formatted_history = []
        if history:
            for msg in history:
                if msg["role"] == "user":
                    formatted_history.append(HumanMessage(content=msg["content"]))
                elif msg["role"] == "assistant":
                    clean_content = re.sub(
                        r"<think>.*?</think>\s*",
                        "",
                        msg["content"],
                        flags=re.DOTALL,
                    )
                    formatted_history.append(AIMessage(content=clean_content.strip()))
                else:
                    formatted_history.append(AIMessage(content=msg["content"]))

        return prompt_template.format_prompt(
            chat_history=formatted_history,
            log_context=log_context_str,
            web_context=web_context_str,
            query=query,
        ).to_messages()

    def query(
        self,
        query: str,
        history: List[Dict] = None,
        use_db_search: bool = True,
        use_web_search: bool = False,
        model_name: Optional[str] = None,
    ):
        log_results = []
        if use_db_search:
            log_results = self.retrieve_logs(query)

        web_results = []
        if use_web_search:
            logger.info(f"[MOCK-QUERY] 联网搜索: {query}")
            web_results = [{"content": "模拟网页结果", "source": "mock.com"}]

        combined_context = {"log_context": log_results, "web_context": web_results}

        for chunk in self.generate_response(
            query,
            combined_context,
            history,
            model_name=model_name,
        ):
            yield chunk